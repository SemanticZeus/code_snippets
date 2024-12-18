from docx import Document
import os

def split_docx(input_file, pages_per_file):
    doc = Document(input_file)
    output_dir = f"{os.path.splitext(input_file)[0]}_split"
    os.makedirs(output_dir, exist_ok=True)

    current_doc = Document()
    page_count = 0
    file_index = 1

    for paragraph in doc.paragraphs:
        current_doc.add_paragraph(paragraph.text)
        page_count += 1

        if page_count >= pages_per_file:
            output_file = os.path.join(output_dir, f"part_{file_index}.docx")
            current_doc.save(output_file)
            print(f"Saved: {output_file}")
            current_doc = Document()
            page_count = 0
            file_index += 1

    if len(current_doc.paragraphs) > 0:  # Save remaining pages
        output_file = os.path.join(output_dir, f"part_{file_index}.docx")
        current_doc.save(output_file)
        print(f"Saved: {output_file}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) != 3:
        print("Usage: python split_docx.py <input_file> <pages_per_file>")
    else:
        input_file = sys.argv[1]
        pages_per_file = int(sys.argv[2])
        split_docx(input_file, pages_per_file)

